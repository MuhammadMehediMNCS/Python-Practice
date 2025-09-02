# Create PDF and Docx File
import os
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4

folder = "DemoFiles"
os.makedirs(folder, exist_ok=True)

word_path = os.path.join(folder, "Students.docx")
pdf_path = os.path.join(folder, "Students.pdf")

# ---------- CREATE ----------
students = ["Mehedi Hasan", "Rahim Uddin", "Karim Mia"]

# Word লিখা
doc = Document()
doc.add_heading("Student List", level=1)
for i, name in enumerate(students, start=1):
    doc.add_paragraph(f"{i}. {name}")
doc.save(word_path)

# PDF লিখা
c = canvas.Canvas(pdf_path, pagesize=A4)
width, height = A4
y = height - 50
c.setFont("Helvetica-Bold", 16)
c.drawString(100, y, "Student List")
y -= 30
c.setFont("Helvetica", 12)
for i, name in enumerate(students, start=1):
    c.drawString(100, y, f"{i}. {name}")
    y -= 20
c.save()

print("✅ CREATE done (Word + PDF).")

# ---------- READ ----------
print("\n📖 Reading data from list:")
for i, name in enumerate(students, start=1):
    print(f"{i}. {name}")

# ---------- UPDATE ----------
students[1] = "Rahim Uddin (Updated)"
print("\n✏️ After update:")
for i, name in enumerate(students, start=1):
    print(f"{i}. {name}")

# ---------- DELETE ----------
del students[0]  # প্রথম জন মুছে ফেলা
print("\n❌ After delete:")
for i, name in enumerate(students, start=1):
    print(f"{i}. {name}")

# ---------- RE-WRITE ----------
# Word আবার লিখা
doc = Document()
doc.add_heading("Student List (Updated)", level=1)
for i, name in enumerate(students, start=1):
    doc.add_paragraph(f"{i}. {name}")
doc.save(word_path)

# PDF আবার লিখা
c = canvas.Canvas(pdf_path, pagesize=A4)
y = height - 50
c.setFont("Helvetica-Bold", 16)
c.drawString(100, y, "Student List (Updated)")
y -= 30
c.setFont("Helvetica", 12)
for i, name in enumerate(students, start=1):
    c.drawString(100, y, f"{i}. {name}")
    y -= 20
c.save()

print("\n✅ Word & PDF re-written after update+delete.")

